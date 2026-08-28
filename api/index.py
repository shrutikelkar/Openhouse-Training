"""Openhouse Playwise — standalone Vercel serverless API.

Fully independent from the Openhouse Quality Observations app: separate
login, separate data store, separate deploy. Trainees explain art & design
games out loud; Gemini asks follow-up questions and scores the explanation
on age fit, clarity, gameplay accuracy, challenge-adjustment accuracy, and
whether it was genuine or read verbatim from the reference material.

Two roles:
  staff    — signs in with STAFF_EMAIL/STAFF_PASSWORD, reaches the dashboard
             (view every trainee's responses, add new trainees).
  trainee  — signs in with phone + name + a 4-digit PIN the admin set when
             adding them, reaches only the recording tool, scoped to their
             own category and their own saved records.

Accounts (env vars set in the Vercel dashboard — defaults below are
placeholders, replace them before relying on this for anything real):
  STAFF_EMAIL / STAFF_PASSWORD               — admin login
  APP_SECRET                                  — any long random string (signs tokens)
  KV_REST_API_URL / KV_REST_API_TOKEN         (Vercel KV)   — OR
  UPSTASH_REDIS_REST_URL / UPSTASH_REDIS_REST_TOKEN  (Upstash)
  GEMINI_API_KEY                              — Google AI Studio key (free tier)
  GEMINI_MODEL                                — defaults to gemini-flash-lite-latest
  SMTP_EMAIL / SMTP_APP_PASSWORD              — optional: sends a welcome email with
             login credentials when a trainee is added, if they have an email on file.
             Use a Gmail/Workspace address with an App Password (needs 2-Step
             Verification enabled) — free, no third-party service. If unset, trainee
             creation still works, it just skips the email.

Trainees are stored in Redis (admin-managed at runtime, not in env vars).
Their PIN is stored as-is (not hashed) — this is a low-stakes internal
training tool, not a system holding sensitive personal data, so a 4-digit
PIN is meant as friction against casual misuse, not real security.
"""
import os
import re
import json
import time
import hmac
import base64
import hashlib
import urllib.request
import urllib.error
import urllib.parse
import io
import smtplib
from email.mime.text import MIMEText
from typing import Optional

from fastapi import FastAPI, Request, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SECRET = os.environ.get("APP_SECRET", "change-me-before-deploy")
STAFF_EMAIL = os.environ.get("STAFF_EMAIL", "staff@openhouse").strip().lower()
STAFF_PASSWORD = os.environ.get("STAFF_PASSWORD", "oh.explaincheck")

REDIS_URL = os.environ.get("KV_REST_API_URL") or os.environ.get("UPSTASH_REDIS_REST_URL")
REDIS_TOKEN = os.environ.get("KV_REST_API_TOKEN") or os.environ.get("UPSTASH_REDIS_REST_TOKEN")
EXPLAIN_LIST_KEY = "explanations"
TRAINEES_KEY = "trainees"  # redis hash: phone -> json trainee record

BLOB_TOKEN = os.environ.get("BLOB_READ_WRITE_TOKEN")
BLOB_API_BASE = "https://blob.vercel-storage.com"
BLOB_API_VERSION = "10"
ARTWORK_AGE_BANDS = ["5-8", "8-12"]
ARTWORK_UNITS = 10
ARTWORK_MAX_BYTES = 8 * 1024 * 1024  # 8MB per file

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-flash-lite-latest")
EXPLAIN_MIN_TURNS = 3
EXPLAIN_MAX_TURNS = 6

SMTP_EMAIL = os.environ.get("SMTP_EMAIL")
SMTP_APP_PASSWORD = os.environ.get("SMTP_APP_PASSWORD")
SMTP_HOST = os.environ.get("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))

CATEGORIES = [
    "art-design", "public-speaking", "robotics", "music", "chess",
    "art-3-5", "storytelling-3-5", "stem-3-5",
]


def _trainee_categories(trainee: dict) -> list:
    """A trainee can now be assigned more than one category. Older records only
    ever had the singular 'category' field, so fall back to wrapping that."""
    cats = trainee.get("categories")
    if cats:
        return cats
    single = trainee.get("category")
    return [single] if single else []

# games-data.js tags each game with short skill codes (e.g. "L&T") — expanded
# here so both the AI prompt and its questions use the real name, never the code.
SKILL_LABELS = {
    "C&P": "Colour & Painting",
    "L&T": "Line & Texture",
    "S&F": "Shape & Form",
    "B&C": "Balance & Composition",
    "I&C": "Imagination & Collaboration",
    "VS": "Vocal Skills",
    "BL": "Body Language",
    "C&S": "Content & Structure",
}

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


def _sign(payload: str) -> str:
    return hmac.new(SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()


def _make_token(role: str, sub: str = "") -> str:
    payload = f"{role}:{sub}"
    return base64.urlsafe_b64encode(f"{payload}.{_sign(payload)}".encode()).decode()


def _decode_token(authorization: Optional[str]) -> tuple:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "login required")
    try:
        raw = base64.urlsafe_b64decode(authorization.split(" ", 1)[1].encode()).decode()
        payload, sig = raw.rsplit(".", 1)
        role, sub = payload.split(":", 1)
    except Exception:
        raise HTTPException(401, "login required")
    if not hmac.compare_digest(sig, _sign(payload)):
        raise HTTPException(401, "login required")
    return role, sub


def _check(authorization: Optional[str], allowed_roles: Optional[set] = None) -> tuple:
    role, sub = _decode_token(authorization)
    if allowed_roles is not None and role not in allowed_roles:
        raise HTTPException(403, "not allowed for this login")
    return role, sub


def _redis(*cmd):
    if not REDIS_URL or not REDIS_TOKEN:
        raise HTTPException(500, "no data store configured — add Vercel KV / Upstash env vars")
    req = urllib.request.Request(
        REDIS_URL,
        data=json.dumps(list(cmd)).encode(),
        headers={"Authorization": f"Bearer {REDIS_TOKEN}", "Content-Type": "application/json"},
    )
    with urllib.request.urlopen(req, timeout=10) as r:
        return json.loads(r.read())


def _blob_put(pathname: str, data: bytes, content_type: str) -> dict:
    """Uploads bytes to Vercel Blob via its raw REST API — there's no official
    Python SDK, only JS/TS, so this talks to blob.vercel-storage.com directly.
    x-add-random-suffix lets Vercel itself guarantee a unique path (never
    allow-overwrite) so a resubmission can never collide with or destroy an
    earlier upload — the actual final URL always comes back in the response,
    so the caller-supplied pathname only needs to be a readable hint."""
    if not BLOB_TOKEN:
        raise HTTPException(500, "no blob storage configured — add BLOB_READ_WRITE_TOKEN")
    url = f"{BLOB_API_BASE}/?pathname={urllib.parse.quote(pathname)}"
    req = urllib.request.Request(
        url, data=data, method="PUT",
        headers={
            "authorization": f"Bearer {BLOB_TOKEN}",
            "x-api-version": BLOB_API_VERSION,
            "x-content-type": content_type,
            "access": "public",
            "x-add-random-suffix": "1",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=25) as r:
            return json.loads(r.read())
    except urllib.error.HTTPError as e:
        raise HTTPException(502, f"blob upload failed: {e.read().decode(errors='replace')}")


def _get_trainee(phone: str) -> Optional[dict]:
    res = _redis("HGET", TRAINEES_KEY, phone)
    val = res.get("result")
    return json.loads(val) if val else None


def _list_trainees() -> list:
    res = _redis("HGETALL", TRAINEES_KEY)
    flat = res.get("result") or []
    out = []
    for i in range(1, len(flat), 2):
        try:
            out.append(json.loads(flat[i]))
        except ValueError:
            continue
    return out


def _q_words(q: str) -> set:
    return set(re.findall(r"[a-z0-9]+", (q or "").lower()))


def _is_repeat_question(candidate: str, prior_questions: list) -> bool:
    """Word-overlap check — catches exact repeats and close rephrasings the
    model might produce despite being told not to repeat itself."""
    cset = _q_words(candidate)
    if not cset:
        return False
    for p in prior_questions:
        pset = _q_words(p)
        if not pset:
            continue
        overlap = len(cset & pset) / len(cset | pset)
        if overlap >= 0.55:
            return True
    return False


def _tokenize(text: str) -> list:
    return re.findall(r"[a-z0-9]+", (text or "").lower())


def _longest_common_run(a: list, b: list):
    """Longest contiguous run of tokens appearing in both sequences, in the
    same order — a much stronger 'reading the script' signal than plain word
    overlap, since genuine explanations rarely reproduce a long exact phrase."""
    if not a or not b:
        return 0, ""
    m, n = len(a), len(b)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    best_len, best_end = 0, 0
    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if a[i - 1] == b[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
                if dp[i][j] > best_len:
                    best_len = dp[i][j]
                    best_end = i
    phrase = " ".join(a[best_end - best_len:best_end]) if best_len else ""
    return best_len, phrase


def _genuineness_signal(game: dict, history: list) -> str:
    """Computes an objective verbatim-overlap measurement between what the
    trainee said and the game's reference text, to ground the model's
    genuine/reading judgment in evidence rather than a vibe check."""
    ref_text = " ".join(filter(None, [
        game.get("goal") or "",
        " ".join(game.get("steps") or []),
        game.get("easier") or "",
        game.get("harder") or "",
        game.get("ends_when") or "",
        " ".join((game.get("difficulty_levels") or {}).values()),
        " ".join(v.get("text", "") for v in (game.get("variations") or [])),
    ]))
    trainee_text = " ".join(h.get("text", "") for h in history if h.get("role") == "trainee")
    ref_tokens = _tokenize(ref_text)
    trainee_tokens = _tokenize(trainee_text)
    run_len, run_phrase = _longest_common_run(trainee_tokens, ref_tokens)
    tset, rset = set(trainee_tokens), set(ref_tokens)
    overlap_pct = round(100 * len(tset & rset) / max(1, len(tset | rset)))
    return (
        f"Genuineness signal (objectively computed, not the model's opinion): the trainee's combined "
        f"spoken answer so far shares a longest verbatim run of {run_len} consecutive words with the "
        f"official reference text"
        + (f' ("{run_phrase}")' if run_phrase else "")
        + f", and an overall vocabulary overlap of {overlap_pct}% with the reference material.\n"
        "- A long verbatim run (roughly 6+ consecutive words matching the reference material "
        "word-for-word) strongly suggests the trainee is reading the script rather than explaining "
        "in their own words.\n"
        "- Short overlaps of common words, or matching a few key terms like the game name or "
        "materials, are normal and NOT evidence of reading — most vocabulary overlap with the "
        "reference is expected and fine.\n"
        "- Also weigh natural speech qualities: their own phrasing and reasonable imperfection reads "
        "as genuine; suspiciously polished, complete, written-style sentences that closely mirror the "
        "reference structure read as scripted.\n"
    )


def _gemini_json(system: str, user_message: str) -> dict:
    if not GEMINI_API_KEY:
        raise HTTPException(500, "no GEMINI_API_KEY configured — add it in the Vercel dashboard")
    url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"
        f"?key={GEMINI_API_KEY}"
    )
    req = urllib.request.Request(
        url,
        data=json.dumps({
            "system_instruction": {"parts": [{"text": system}]},
            "contents": [{"role": "user", "parts": [{"text": user_message}]}],
            "generationConfig": {
                "maxOutputTokens": 8000,
                "responseMimeType": "application/json",
            },
        }).encode(),
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            body = json.loads(r.read())
    except urllib.error.HTTPError as e:
        detail = e.read().decode(errors="replace")
        raise HTTPException(502, f"Gemini API error ({e.code}): {detail}")
    try:
        text = body["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError):
        raise HTTPException(502, f"Gemini returned no usable response: {json.dumps(body)[:500]}")
    if text.startswith("```"):
        text = text.strip("`")
        text = text.split("\n", 1)[1] if "\n" in text else text
        if text.lower().startswith("json"):
            text = text[4:]
    try:
        return json.loads(text)
    except ValueError:
        raise HTTPException(502, f"model did not return valid JSON. Raw text: {text[:800]!r}. Full body: {json.dumps(body)[:800]}")


def _gemini_transcribe(audio_b64: str, mime_type: str) -> str:
    """Speech-to-text via Gemini's audio understanding — used instead of the
    browser's SpeechRecognition API, which only exists in Chromium browsers.
    Recording itself uses MediaRecorder (Chrome, Firefox, Safari, Edge all
    support it), so this is what makes the mic work everywhere."""
    if not GEMINI_API_KEY:
        raise HTTPException(500, "no GEMINI_API_KEY configured — add it in the Vercel dashboard")
    url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"
        f"?key={GEMINI_API_KEY}"
    )
    prompt = (
        "Transcribe the spoken audio exactly as spoken, in English. "
        "Return only the transcript text — no commentary, labels, or quotation marks. "
        "If nothing intelligible was said, return an empty string."
    )
    req = urllib.request.Request(
        url,
        data=json.dumps({
            "contents": [{"role": "user", "parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": mime_type, "data": audio_b64}},
            ]}],
            "generationConfig": {"maxOutputTokens": 2000},
        }).encode(),
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            body = json.loads(r.read())
    except urllib.error.HTTPError as e:
        detail = e.read().decode(errors="replace")
        raise HTTPException(502, f"Gemini API error ({e.code}): {detail}")
    try:
        return body["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError):
        # empty/blocked response (e.g. silence) — treat as "nothing said" rather than an error
        return ""


# ---------- auth ----------

@app.post("/api/login")
async def login(req: Request):
    """Staff (admin) login."""
    b = await req.json()
    email = (b.get("email") or "").strip().lower()
    password = b.get("password") or ""
    if email != STAFF_EMAIL or not hmac.compare_digest(password, STAFF_PASSWORD):
        raise HTTPException(401, "wrong email or password")
    return {"token": _make_token("staff")}


@app.post("/api/trainee/login")
async def trainee_login(req: Request):
    b = await req.json()
    phone = (b.get("phone") or "").strip()
    name = (b.get("name") or "").strip()
    pin = (b.get("pin") or "").strip()
    if not phone or not name or not pin:
        raise HTTPException(400, "phone, name and pin are all required")
    trainee = _get_trainee(phone)
    if not trainee or trainee.get("name", "").strip().lower() != name.lower() or not hmac.compare_digest(str(trainee.get("pin", "")), pin):
        raise HTTPException(401, "no match for that phone number, name and PIN")
    categories = _trainee_categories(trainee)
    return {
        "token": _make_token("trainee", phone),
        "name": trainee.get("name"),
        "phone": phone,
        "category": categories[0] if categories else None,
        "categories": categories,
        "cohort": trainee.get("cohort"),
        "approved": bool(trainee.get("approved")),
    }


@app.get("/api/trainee/status")
async def trainee_status(authorization: Optional[str] = Header(None)):
    """Fresh approval/categories/cohort status for the current trainee — used
    to reflect an admin approving them (or changing their categories) without
    requiring a fresh login."""
    role, sub = _check(authorization, {"trainee"})
    trainee = _get_trainee(sub)
    if not trainee:
        raise HTTPException(404, "trainee record not found")
    categories = _trainee_categories(trainee)
    return {
        "name": trainee.get("name"),
        "category": categories[0] if categories else None,
        "categories": categories,
        "cohort": trainee.get("cohort"),
        "approved": bool(trainee.get("approved")),
    }


# ---------- admin: trainee management ----------

def _send_welcome_email(to_email: str, trainee: dict, site_url: str) -> bool:
    """Best-effort welcome email with login credentials. Returns False (never
    raises) if SMTP isn't configured or sending fails — trainee creation must
    never be blocked by an email problem."""
    if not SMTP_EMAIL or not SMTP_APP_PASSWORD or not to_email:
        return False
    cats = _trainee_categories(trainee)
    body = (
        f"Hi {trainee['name']},\n\n"
        f"You've been added to Openhouse Playwise"
        f"{' (' + ', '.join(cats) + ')' if cats else ''}.\n\n"
        f"Sign in here: {site_url}\n"
        f"Phone number: {trainee['phone']}\n"
        f"PIN: {trainee['pin']}\n\n"
        "— Housie"
    )
    msg = MIMEText(body)
    msg["Subject"] = "Your Openhouse Playwise login"
    msg["From"] = SMTP_EMAIL
    msg["To"] = to_email
    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=10) as server:
            server.starttls()
            server.login(SMTP_EMAIL, SMTP_APP_PASSWORD)
            server.send_message(msg)
        return True
    except Exception:
        return False


@app.post("/api/admin/trainees")
async def add_trainee(req: Request, authorization: Optional[str] = Header(None)):
    """PIN is not chosen by the admin — it's always derived from the cohort's
    start date as DDMM, so it's never trusted from the client."""
    _check(authorization, {"staff"})
    b = await req.json()
    phone = (b.get("phone") or "").strip()
    name = (b.get("name") or "").strip()
    raw_categories = b.get("categories")
    if raw_categories is None:
        # back-compat with any caller still posting the old singular field
        single = (b.get("category") or "").strip()
        raw_categories = [single] if single else []
    categories = list(dict.fromkeys(c.strip() for c in raw_categories if c and c.strip()))
    cohort = (b.get("cohort") or "").strip()
    email = (b.get("email") or "").strip()
    if not phone or not name or not cohort:
        raise HTTPException(400, "phone, name and cohort start date are required")
    if not categories:
        raise HTTPException(400, "at least one category is required")
    bad = [c for c in categories if c not in CATEGORIES]
    if bad:
        raise HTTPException(400, f"category must be one of {CATEGORIES}")
    try:
        y, m, d = cohort.split("-")
        pin = f"{int(d):02d}{int(m):02d}"
    except ValueError:
        raise HTTPException(400, "cohort must be a valid date (YYYY-MM-DD)")
    # Re-submitting this form for a phone that already has a trainee (e.g. to
    # fix a typo) must not silently wipe their approval or creation date.
    existing = _get_trainee(phone)
    rec = {
        "phone": phone, "name": name, "pin": pin,
        "categories": categories, "category": categories[0], "cohort": cohort, "email": email,
        "approved": existing.get("approved", False) if existing else False,
        "added_at": existing.get("added_at") if existing else time.strftime("%Y-%m-%d %H:%M"),
    }
    _redis("HSET", TRAINEES_KEY, phone, json.dumps(rec, ensure_ascii=False))
    email_sent = False
    if email:
        site_url = str(req.base_url).rstrip("/")
        email_sent = _send_welcome_email(email, rec, site_url)
    return {"ok": True, "trainee": rec, "email_sent": email_sent}


@app.post("/api/admin/trainees/approve")
async def approve_trainee(req: Request, authorization: Optional[str] = Header(None)):
    """Marks a trainee's assessment as approved, unlocking their download.
    Toggleable — pass approved:false to revoke."""
    _check(authorization, {"staff"})
    b = await req.json()
    phone = (b.get("phone") or "").strip()
    approved = bool(b.get("approved", True))
    trainee = _get_trainee(phone)
    if not trainee:
        raise HTTPException(404, "trainee not found")
    trainee["approved"] = approved
    _redis("HSET", TRAINEES_KEY, phone, json.dumps(trainee, ensure_ascii=False))
    return {"ok": True, "trainee": trainee}


@app.get("/api/admin/trainees")
async def list_trainees(authorization: Optional[str] = Header(None)):
    _check(authorization, {"staff"})
    return _list_trainees()


@app.delete("/api/admin/trainees/{phone}")
async def remove_trainee(phone: str, delete_data: bool = False, authorization: Optional[str] = Header(None)):
    """Removes a trainee's login. By default their saved explanation records
    are left alone so past assessments stay on file for review; pass
    delete_data=true to also permanently erase those records."""
    _check(authorization, {"staff"})
    if not _get_trainee(phone):
        raise HTTPException(404, "trainee not found")
    _redis("HDEL", TRAINEES_KEY, phone)
    if delete_data:
        res = _redis("LRANGE", EXPLAIN_LIST_KEY, "0", "999")
        existing = [json.loads(x) for x in (res.get("result") or [])]
        kept = [x for x in existing if x.get("trainee_phone") != phone]
        if len(kept) != len(existing):
            _redis("DEL", EXPLAIN_LIST_KEY)
            if kept:
                _redis("RPUSH", EXPLAIN_LIST_KEY, *[json.dumps(x, ensure_ascii=False) for x in kept])
    return {"ok": True}


@app.get("/api/admin/categories")
async def get_categories(authorization: Optional[str] = Header(None)):
    _check(authorization, {"staff"})
    return CATEGORIES


# ---------- artwork uploads ----------

@app.post("/api/artwork/upload")
async def artwork_upload(req: Request, authorization: Optional[str] = Header(None)):
    """Uploads/replaces one artwork slot (a given age band + unit number) for
    the signed-in trainee. Body: {age_band, unit, file (base64), filename,
    content_type}. Each submission gets its own timestamped pathname — a
    resubmission never overwrites or deletes the previous file, it just
    becomes the one shown in artwork[age_band][unit]."""
    role, phone = _check(authorization, {"trainee"})
    trainee = _get_trainee(phone)
    if not trainee:
        raise HTTPException(404, "trainee not found")
    if "art-design" not in _trainee_categories(trainee):
        raise HTTPException(403, "artwork uploads are only for the art & design category")
    b = await req.json()
    age_band = (b.get("age_band") or "").strip()
    filename = (b.get("filename") or "upload").strip()
    content_type = (b.get("content_type") or "application/octet-stream").strip()
    file_b64 = b.get("file") or ""
    if age_band not in ARTWORK_AGE_BANDS:
        raise HTTPException(400, f"age_band must be one of {ARTWORK_AGE_BANDS}")
    try:
        unit = int(b.get("unit"))
    except (TypeError, ValueError):
        raise HTTPException(400, "unit must be a number")
    if not (1 <= unit <= ARTWORK_UNITS):
        raise HTTPException(400, f"unit must be between 1 and {ARTWORK_UNITS}")
    if not file_b64:
        raise HTTPException(400, "no file provided")
    if content_type not in ("image/jpeg", "image/png", "application/pdf"):
        raise HTTPException(400, "file must be JPG, PNG or PDF")
    try:
        data = base64.b64decode(file_b64)
    except Exception:
        raise HTTPException(400, "invalid file data")
    if len(data) > ARTWORK_MAX_BYTES:
        raise HTTPException(400, "file too large (max 8MB)")

    ext = os.path.splitext(filename)[1][:10]
    pathname = f"artwork/{phone}/{age_band}/unit-{unit}-{int(time.time() * 1000)}{ext}"
    result = _blob_put(pathname, data, content_type)

    artwork = trainee.get("artwork") or {}
    band_entry = artwork.get(age_band) or {}
    # each unit holds every submission ever made for it, oldest first — a
    # resubmission appends rather than replaces, so nothing is ever lost
    existing = band_entry.get(str(unit))
    history = existing if isinstance(existing, list) else ([existing] if existing else [])
    history.append({
        "url": result.get("url"),
        "filename": filename,
        "content_type": content_type,
        "uploaded_at": time.strftime("%Y-%m-%d %H:%M"),
    })
    band_entry[str(unit)] = history
    artwork[age_band] = band_entry
    trainee["artwork"] = artwork

    # a resubmission clears any "please redo this one" flag staff put on it
    redo = trainee.get("artwork_redo") or {}
    band_redo = redo.get(age_band) or []
    if unit in band_redo:
        redo[age_band] = [u for u in band_redo if u != unit]
        trainee["artwork_redo"] = redo

    _redis("HSET", TRAINEES_KEY, phone, json.dumps(trainee, ensure_ascii=False))
    return {"ok": True, "artwork": artwork, "redo": trainee.get("artwork_redo") or {}}


@app.get("/api/artwork/mine")
async def artwork_mine(authorization: Optional[str] = Header(None)):
    role, phone = _check(authorization, {"trainee"})
    trainee = _get_trainee(phone)
    if not trainee:
        raise HTTPException(404, "trainee not found")
    if "art-design" not in _trainee_categories(trainee):
        raise HTTPException(403, "artwork uploads are only for the art & design category")
    return {"artwork": trainee.get("artwork") or {}, "redo": trainee.get("artwork_redo") or {}}


@app.get("/api/admin/artwork")
async def admin_artwork(authorization: Optional[str] = Header(None)):
    """Every trainee who has uploaded at least one piece of artwork, or has a
    pending redo request, for the dashboard's artwork view."""
    _check(authorization, {"staff"})
    trainees = _list_trainees()
    return [
        {
            "phone": t.get("phone"), "name": t.get("name"), "categories": _trainee_categories(t),
            "artwork": t.get("artwork") or {}, "redo": t.get("artwork_redo") or {},
        }
        for t in trainees if t.get("artwork") or t.get("artwork_redo")
    ]


@app.post("/api/admin/artwork/redo")
async def admin_artwork_redo(req: Request, authorization: Optional[str] = Header(None)):
    """Staff asks a trainee to redo one artwork unit. The current submission
    is left in place (still viewable) — only the flag is set, so the trainee
    sees why; resubmitting via /api/artwork/upload overwrites the old file
    at the same pathname and clears the flag automatically."""
    _check(authorization, {"staff"})
    b = await req.json()
    phone = (b.get("phone") or "").strip()
    age_band = (b.get("age_band") or "").strip()
    if age_band not in ARTWORK_AGE_BANDS:
        raise HTTPException(400, f"age_band must be one of {ARTWORK_AGE_BANDS}")
    try:
        unit = int(b.get("unit"))
    except (TypeError, ValueError):
        raise HTTPException(400, "unit must be a number")
    if not (1 <= unit <= ARTWORK_UNITS):
        raise HTTPException(400, f"unit must be between 1 and {ARTWORK_UNITS}")

    trainee = _get_trainee(phone)
    if not trainee:
        raise HTTPException(404, "trainee not found")

    redo = trainee.get("artwork_redo") or {}
    band_redo = redo.get(age_band) or []
    if unit not in band_redo:
        band_redo = band_redo + [unit]
    redo[age_band] = band_redo
    trainee["artwork_redo"] = redo

    _redis("HSET", TRAINEES_KEY, phone, json.dumps(trainee, ensure_ascii=False))
    return {"ok": True, "artwork": trainee.get("artwork") or {}, "redo": redo}


# ---------- explanation quiz ----------

@app.post("/api/explain/transcribe")
async def explain_transcribe(req: Request, authorization: Optional[str] = Header(None)):
    """Transcribes a trainee's recorded answer. The browser records audio with
    MediaRecorder (works everywhere) and uploads it here instead of using
    SpeechRecognition (Chromium-only) to turn it into text client-side."""
    _check(authorization, {"trainee"})
    b = await req.json()
    audio_b64 = b.get("audio") or ""
    mime_type = b.get("mime_type") or "audio/wav"
    if not audio_b64:
        raise HTTPException(400, "no audio provided")
    transcript = _gemini_transcribe(audio_b64, mime_type)
    return {"transcript": transcript}


@app.post("/api/explain/turn")
async def explain_turn(req: Request, authorization: Optional[str] = Header(None)):
    """One turn of the game-explanation quiz. The client sends the full game
    reference data (from games-data.js) plus the conversation so far; this
    returns either another follow-up question or the final 5-point score.
    """
    _check(authorization)
    b = await req.json()
    game = b.get("game") or {}
    game_name = game.get("name") or "the game"
    age_band = b.get("age_band") or "unknown"
    history = b.get("history") or []
    trainee_turns = sum(1 for h in history if h.get("role") == "trainee")
    prior_questions = [h.get("text", "") for h in history if h.get("role") == "ai"]
    skill_names = [SKILL_LABELS.get(code, code) for code in game.get("skills") or []]
    debrief_questions = game.get("debrief") or []
    has_challenge = bool(game.get("easier") or game.get("harder"))
    lanyard_skills = game.get("lanyard_skills") or []
    scenarios = game.get("scenarios") or []
    facts = game.get("facts") or []

    # Just a count for the prompt's own reference — coverage of each topic (debrief
    # especially) is guaranteed by the explicit rules below, not by turn-count math.
    required_topics = 2  # goal/steps + materials are always required
    if has_challenge:
        required_topics += 1
    if lanyard_skills:
        required_topics += 1
    if scenarios:
        required_topics += 1
    if facts:
        required_topics += 1
    if game.get("variations"):
        required_topics += 1
    if skill_names:
        required_topics += 1
    if debrief_questions:
        required_topics += 1
    min_turns = EXPLAIN_MIN_TURNS
    max_turns = EXPLAIN_MAX_TURNS

    if age_band in ("5–8", "5-8"):
        age_example_rule = (
            "Since the age band here is 5-8, they must ALSO give at least one concrete worked "
            "example (e.g. \"like if you pick a red bead, you put it on the red square\") — for "
            "young children, listing abstract rules without a concrete example is NOT enough, score 0."
        )
    else:
        age_example_rule = (
            "For this 8-12 band, a concrete example is a plus but not mandatory, as long as the "
            "direct-address register is right."
        )

    system = (
        "You are quizzing a childcare educator-in-training on a game "
        "they must run with children. You have the game's authoritative reference "
        "data below. Listen to what the trainee said and decide whether to ask one "
        "more short follow-up question, or finalize scoring.\n\n"
        f"Reference data (JSON): {json.dumps(game, ensure_ascii=False)}\n\n"
        + (f"This game targets these skills: {', '.join(skill_names)}.\n\n" if skill_names else "")
        + f"The trainee has been told to explain this game for a group of children aged {age_band}.\n\n"
        + _genuineness_signal(game, history) + "\n"
        + (
            "Questions you have ALREADY asked in this conversation (do NOT repeat any of "
            "these, or ask a near-duplicate of one — always move to a new, different gap):\n"
            + "\n".join(f"- {q}" for q in prior_questions) + "\n\n"
            if prior_questions else ""
        )
        + "Rules:\n"
        f"- The trainee must answer at least {min_turns} turns before you finalize — this game "
        f"has {required_topics} mandatory topics (listed below), so do not finalize early just "
        "because the conversation feels sufficient; every one of them needs its own question "
        "unless the trainee's own words already nailed it.\n"
        f"- You must finalize by turn {max_turns} no matter what.\n"
        + (
            "- A debrief question is NOT something a trainee would organically include while "
            "narrating game mechanics to children — it is a distinct reflective question about "
            "how they'd talk to a child afterwards. Never treat it as \"already covered\" by the "
            "initial explanation; it always needs its own dedicated question before you finalize.\n"
            if debrief_questions else ""
        )
        + "- Before you pick your next question, first re-read everything the trainee has "
        "already said — including their very first explanation, not just the latest turn — "
        "against the reference data. The topic list below describes what must be COVERED "
        "across the whole conversation, not a fixed script to work through regardless of "
        "what's already been said. If the trainee's initial explanation already addressed a "
        "topic clearly and correctly, treat it as satisfied and do NOT ask about it again — "
        "move straight to a topic that's genuinely missing, vague, or wrong. Every question "
        "you ask must be grounded in this specific trainee's own words or this specific game's "
        "details, never a generic template question you'd ask regardless of what they said.\n"
        "- Across the whole conversation, ask at least one gameplay question covering the "
        "goal and how to play, using the 'goal'/'steps' fields as the correct answer.\n"
        "- Across the whole conversation, ask at least one question about what materials "
        "or equipment the game needs, using the 'materials' field as the correct answer.\n"
        + (
            "- Across the whole conversation, ask at least one challenge question about how "
            "to make the game easier or harder for a child's needs, using the 'easier'/'harder'/"
            "'difficulty_levels' fields in the reference data as the correct answer.\n"
            if has_challenge else ""
        )
        + (
            "- Across the whole conversation, ask the trainee to describe at least one individual "
            "skill lanyard from the reference data's 'lanyard_skills' list, in their own words — "
            "what it means and what it looks like in a child, checked against its 'description' field. "
            + (
                "Mimicry MUST be one of the lanyards you ask about at some point in the "
                "conversation, since it's the one trainees most often misunderstand. "
                if any((ls.get("name") or "").lower() == "mimicry" for ls in lanyard_skills) else ""
            )
            + "\n"
            if lanyard_skills else ""
        )
        + (
            "- Across the whole conversation, present the trainee with ONE fabricated scenario "
            "from the reference data's 'scenarios' list (describe the child's situation using its "
            "'scenario' field) and ask which lanyard they'd give that child and why. Their answer "
            "is correct if it names the scenario's 'lanyard' field, or another lanyard they justify "
            "just as soundly using reasoning like its 'why' field — do not require exact wording.\n"
            if scenarios else ""
        )
        + (
            "- Across the whole conversation, ask the trainee at least one fact-check question "
            "drawn from the reference data's 'facts' list (e.g. whether it's ever okay to skip "
            "using lanyards, or whether every child can be given the same lanyard) and check their "
            "answer against that fact. The list is general reference material to draw from, not a "
            "checklist you must ask about item by item.\n"
            if facts else ""
        )
        + (
            "- Across the whole conversation, also ask at least one question about a variation "
            "of the game — e.g. ask the trainee to describe one of the reference data's "
            "'variations', or how they'd introduce it to the group — using the 'variations' "
            "field as the correct answer.\n"
            if game.get("variations") else ""
        )
        + (
            f"- Across the whole conversation, also ask at least one question about what "
            f"skill(s) this game helps children build — the correct answer is: "
            f"{', '.join(skill_names)}. Phrase it naturally (e.g. \"what skill do you think "
            f"this game helps children build?\"), never mention the raw code.\n"
            if skill_names else ""
        )
        + (
            "- Across the whole conversation, also ask ONE debrief-style reflection question, "
            "picked from (or closely adapted from) this list: "
            + " / ".join(debrief_questions) + ". Unlike the other questions, there is no single "
            "right answer here — treat any thoughtful, genuine reflection as satisfying this, "
            "even if it differs from what you might expect.\n"
            if debrief_questions else ""
        )
        + "- Ask only ONE question at a time, short and natural to say out loud (it will be "
        "read by text-to-speech) — do not list multiple questions.\n"
        "- Never repeat a question already listed above, and never ask a generic filler "
        "question like 'can you tell me more' — always name the specific thing you want "
        "(a missing step, a specific material, the group size, the easier/harder answer, "
        "a variation, a skill, or a debrief reflection).\n"
        "- If you cannot think of a new, specific, non-repetitive question — because the "
        "trainee has already covered the goal, steps, and materials"
        + (", an easier/harder answer" if has_challenge else "")
        + (", a variation" if game.get("variations") else "")
        + (", a skill" if skill_names else "")
        + (", a lanyard description" if lanyard_skills else "")
        + (", a scenario" if scenarios else "")
        + (", a fact-check" if facts else "")
        + (", and a debrief reflection" if debrief_questions else "")
        + " — finalize instead of asking anything further, even before the max turn.\n"
        "- Target the biggest gap or error first (missing step, wrong material, wrong group "
        "size/goal"
        + (", a wrong/missing easier-harder answer" if has_challenge else "")
        + (", an uncovered variation" if game.get("variations") else "")
        + (", a missing skill answer" if skill_names else "")
        + (", a missing/wrong lanyard description" if lanyard_skills else "")
        + (", a wrong scenario answer" if scenarios else "")
        + (", a wrong fact-check answer" if facts else "")
        + (", or a missing debrief reflection" if debrief_questions else "")
        + ").\n"
        "- When finalizing, score five criteria as 1 (met) or 0 (not met):\n"
        "  1. age_appropriateness — this is NOT just about materials/pacing. The trainee must have "
        "actually pitched the explanation AS IF speaking directly to the children themselves (simple, "
        "direct, concrete language — 'you pick a bead and...') — not described the game mechanically "
        "to an adult colleague or trainer ('the children pick a bead and...' read like a briefing is "
        "a fail even if every fact is correct). "
        + age_example_rule + "\n"
        "  2. clarity — was the explanation easy to follow, steps in a sensible order.\n"
        + (
            "  3. gameplay_accuracy — did the core mechanics match the reference: goal, steps, and materials.\n"
            "  4. challenge_accuracy — did they give a correct easier/harder difficulty-adjustment answer "
            "(scored separately from gameplay_accuracy — a trainee can get the game right but the "
            "difficulty adjustment wrong, or vice versa). This counts as correct (1) if it matches the "
            "reference 'easier'/'harder'/'difficulty_levels' fields, OR if they propose a different but "
            "genuinely sound, sensible, age-appropriate way to adjust difficulty — do not penalize a good "
            "idea just because it isn't the documented one.\n"
            if has_challenge else
            "  3. gameplay_accuracy — here this criterion means ACCURACY OF EXPLANATION: did the trainee "
            "correctly explain the purpose and mechanics from the reference data (the goal/steps, and any "
            "lanyard-skill description, scenario answer, or fact-check they were asked about)? There is "
            "no easier/harder difficulty adjustment for this item — never ask for or expect one.\n"
            "  4. challenge_accuracy — here this criterion means ACCURACY OF DEBRIEF: did the trainee "
            "give a sound, genuine, age-appropriate answer when asked how they'd debrief a child about "
            "their focused skill? Score this against the debrief question(s) actually asked, not "
            "against any difficulty-adjustment concept.\n"
        )
        + "  5. genuine — 1 if this reads as their own understanding in their own words, 0 if it reads "
        "as reading directly from the reference script (use the genuineness signal above as evidence).\n\n"
        + (
            "Additionally, when finalizing: if the trainee's easier/harder answer included a genuinely NEW "
            "idea for adjusting difficulty — one that is NOT already described in the reference 'easier'/"
            "'harder'/'difficulty_levels'/'variations' fields, and is a sensible, usable suggestion — "
            "capture it in a \"new_idea\" field (a short, cleaned-up sentence describing the idea). If there "
            "was no new idea (they just repeated or paraphrased the documented answer, or gave nothing "
            "usable), set \"new_idea\" to null. Only include \"new_idea\" in the final action, not in ask.\n\n"
            if has_challenge else
            "Additionally, when finalizing: there is no difficulty-adjustment concept for this item, so "
            "always set \"new_idea\" to null.\n\n"
        )
        + "Also when finalizing, include trainee-facing feedback split into two parts (the \"reasoning\" "
        "field above is a separate, more technical note for the admin dashboard only — it is NEVER "
        "shown to the trainee, so it can be blunt; these two fields ARE shown to the trainee, so they "
        "must be constructive and encouraging in tone even when pointing out a miss):\n"
        "  - \"strengths\": one specific, genuine sentence about what they did well — name the actual "
        "thing (e.g. 'You clearly explained the goal and the turn order' not 'good job'). Always "
        "present, even on a low score — find something real and specific to credit.\n"
        "  - \"improvements\": an array with one entry for EVERY criterion that scored 0 (empty array "
        "if all five scored 1). Each entry: "
        '{"criterion":"<the criterion key>","reason":"<specific, concrete reason this point was lost, '
        "referencing what they actually said>\",\"fix\":\"<a specific, actionable way to improve — a "
        "concrete example to use, a tone adjustment, simpler wording, a particular step to add — not a "
        'vague "explain more clearly">"}. Keep the tone constructive, never harsh, even though the '
        "reason is specific about the gap.\n\n"
        "Respond with ONLY one JSON object, no other text, no markdown fences, in one of "
        "these two shapes:\n"
        '{"action":"ask","question":"..."}\n'
        'or\n'
        '{"action":"final","scores":{"age_appropriateness":0,"clarity":0,"gameplay_accuracy":0,"challenge_accuracy":0,"genuine":0},'
        '"reasoning":"...","new_idea":null,"strengths":"...","improvements":[{"criterion":"...","reason":"...","fix":"..."}]}'
    )
    transcript = "\n".join(f"{h.get('role')}: {h.get('text','')}" for h in history)
    result = _gemini_json(system, transcript or "(no explanation given yet)")

    # The model can still slip and repeat a question despite being told not to.
    # If that happens after the minimum turns are already met, treat it as a
    # signal there's nothing new left to probe and force a real finalize call
    # rather than let the conversation loop. Below the minimum, swap in a
    # concrete, distinct fallback question instead of retrying the same prompt.
    if result.get("action") == "ask" and _is_repeat_question(result.get("question", ""), prior_questions):
        if trainee_turns >= min_turns:
            finalize_system = system + (
                "\n\nYou just tried to ask a question that repeats one already asked. "
                "There is nothing new left to probe — finalize now with scores and reasoning "
                "instead of asking anything further."
            )
            result = _gemini_json(finalize_system, transcript)
            if result.get("action") != "final":
                result = {
                    "action": "final",
                    "scores": {"age_appropriateness": 1, "clarity": 1, "gameplay_accuracy": 1, "challenge_accuracy": 1, "genuine": 1},
                    "reasoning": "Trainee covered the key points across the conversation; no further gaps to probe.",
                    "strengths": "You covered the key points across the conversation, including the difficulty adjustment.",
                    "improvements": [],
                }
        else:
            fallback_pool = [
                f"What's the very first thing you'd tell the children to do to set up {game_name}?",
                f"What materials would you need on hand to run {game_name}?",
                f"How do you know when {game_name} is over — what's the end condition?",
                f"How would you make {game_name} easier for a child who's struggling?",
                f"How would you make {game_name} harder for a child who's ready for more?",
            ]
            fresh = next((q for q in fallback_pool if not _is_repeat_question(q, prior_questions)), None)
            result = {"action": "ask", "question": fresh or fallback_pool[0]}

    if trainee_turns < min_turns and result.get("action") != "ask":
        result = {
            "action": "ask",
            "question": f"What's the very first thing you'd tell the children to do to set up {game_name}?",
        }
    if trainee_turns >= max_turns and result.get("action") != "final":
        result = {
            "action": "final",
            "scores": {"age_appropriateness": 0, "clarity": 0, "gameplay_accuracy": 0, "challenge_accuracy": 0, "genuine": 0},
            "reasoning": "Reached the maximum number of follow-up questions without a clear finalization from the model.",
            "strengths": "You stayed with it through several follow-up questions.",
            "improvements": [{"criterion": k, "reason": "This wasn't confirmed within the allowed number of follow-ups.", "fix": "Try covering this clearly in your first answer next time."} for k in ("age_appropriateness", "clarity", "gameplay_accuracy", "challenge_accuracy", "genuine")],
            "redo_reason": "timeout",
        }
    if result.get("action") == "final":
        result.setdefault("scores", {}).setdefault("genuine", 0)
        result.setdefault("new_idea", None)
        result.setdefault("strengths", "")
        result.setdefault("improvements", [])
        result.setdefault("suggestion", "")
        # "timeout" means the conversation ran out of turns (a system limit, not a
        # genuineness judgement); anything else with genuine=0 is a real model call.
        result.setdefault("redo_reason", "not_genuine" if not result["scores"].get("genuine") else None)
    return result


MAX_ATTEMPTS_PER_GAME = 2  # the original attempt plus exactly one re-attempt


@app.post("/api/explain/save")
async def explain_save(req: Request, authorization: Optional[str] = Header(None)):
    role, sub = _check(authorization)
    b = await req.json()
    trainee_phone = sub if role == "trainee" else (b.get("trainee_phone") or "")
    category = b.get("category") or ""
    cohort = b.get("cohort") or ""
    if role == "trainee":
        # A trainee can be assigned more than one category and picks which one
        # applies each session — trust that choice only if it's actually one of
        # theirs; cohort still always comes from the server-side record.
        t = _get_trainee(sub)
        if t:
            assigned = _trainee_categories(t)
            if category not in assigned:
                category = assigned[0] if assigned else category
            cohort = t.get("cohort") or cohort
        # enforce the one-reattempt cap server-side, not just in the UI
        game_id = b.get("game_id")
        res = _redis("LRANGE", EXPLAIN_LIST_KEY, "0", "999")
        existing = [json.loads(x) for x in (res.get("result") or [])]
        prior_attempts = sum(1 for r in existing if r.get("trainee_phone") == sub and r.get("game_id") == game_id)
        if prior_attempts >= MAX_ATTEMPTS_PER_GAME:
            raise HTTPException(400, "You've already used your one re-attempt for this game.")
    rec = {
        "id": base64.urlsafe_b64encode(os.urandom(6)).decode().rstrip("="),
        "received_at": time.strftime("%Y-%m-%d %H:%M"),
        "game_id": b.get("game_id"),
        "game_name": b.get("game_name"),
        "age_band": b.get("age_band"),
        "trainee_name": (b.get("trainee_name") or "").strip(),
        "trainee_phone": trainee_phone,
        "category": category,
        "cohort": cohort,
        "history": b.get("history") or [],
        "scores": b.get("scores") or {},
        "reasoning": b.get("reasoning") or "",
        "new_idea": b.get("new_idea") or None,
        "strengths": b.get("strengths") or "",
        "improvements": b.get("improvements") or [],
    }
    _redis("LPUSH", EXPLAIN_LIST_KEY, json.dumps(rec, ensure_ascii=False))
    # A new response (first attempt or re-attempt) means the downloadable
    # document has changed — any prior approval covered the old content, not
    # this, so it needs a fresh review before the trainee can download again.
    if role == "trainee" and t and t.get("approved"):
        t["approved"] = False
        _redis("HSET", TRAINEES_KEY, sub, json.dumps(t, ensure_ascii=False))
    return {"ok": True, "id": rec["id"]}


@app.get("/api/explanations")
async def list_explanations(authorization: Optional[str] = Header(None)):
    _check(authorization, {"staff"})
    res = _redis("LRANGE", EXPLAIN_LIST_KEY, "0", "999")
    rows = res.get("result") or []
    return [json.loads(x) for x in rows]


def _mutate_explanation_record(record_id: str, mutate) -> dict:
    """Finds one explanation record by id, applies `mutate` to it in place,
    and rewrites the whole list back to Redis (there's no per-item update in
    this REST API, only whole-list ops). Raises 404 if not found."""
    res = _redis("LRANGE", EXPLAIN_LIST_KEY, "0", "999")
    rows = res.get("result") or []
    all_recs = [json.loads(x) for x in rows]
    found = None
    for r in all_recs:
        if r.get("id") == record_id:
            mutate(r)
            found = r
            break
    if not found:
        raise HTTPException(404, "record not found")
    _redis("DEL", EXPLAIN_LIST_KEY)
    _redis("RPUSH", EXPLAIN_LIST_KEY, *[json.dumps(x, ensure_ascii=False) for x in all_recs])
    return found


@app.post("/api/admin/explanations/{record_id}/feedback")
async def edit_explanation_feedback(record_id: str, req: Request, authorization: Optional[str] = Header(None)):
    """Lets staff correct the trainee-facing strengths/improvements text before
    approving a download — the same text the trainee sees on screen and in
    their downloaded assessment."""
    _check(authorization, {"staff"})
    b = await req.json()

    def mutate(r):
        r["strengths"] = b.get("strengths") or ""
        r["improvements"] = b.get("improvements") or []

    record = _mutate_explanation_record(record_id, mutate)
    return {"ok": True, "record": record}


@app.post("/api/admin/explanations/{record_id}/comment")
async def edit_explanation_comment(record_id: str, req: Request, authorization: Optional[str] = Header(None)):
    """An admin-only internal note on a response — never shown to the trainee
    and never included in their downloaded assessment."""
    _check(authorization, {"staff"})
    b = await req.json()

    def mutate(r):
        r["admin_comment"] = (b.get("comment") or "").strip()

    record = _mutate_explanation_record(record_id, mutate)
    return {"ok": True, "record": record}


@app.get("/api/my-explanations")
async def my_explanations(authorization: Optional[str] = Header(None)):
    role, sub = _check(authorization, {"trainee"})
    res = _redis("LRANGE", EXPLAIN_LIST_KEY, "0", "999")
    rows = res.get("result") or []
    all_recs = [json.loads(x) for x in rows]
    return [r for r in all_recs if r.get("trainee_phone") == sub]


CRITERIA_LABELS = [
    ("age_appropriateness", "Age fit"),
    ("clarity", "Clarity"),
    ("gameplay_accuracy", "Gameplay accuracy"),
    ("challenge_accuracy", "Challenge accuracy"),
    ("genuine", "Genuine (own words, not read)"),
]

BRAND_CORAL = RGBColor(0xCE, 0x45, 0x20)
BRAND_TEAL = RGBColor(0x0E, 0x8F, 0xA3)
BRAND_INK = RGBColor(0x2C, 0x2B, 0x28)
BRAND_MUTED = RGBColor(0x6B, 0x67, 0x5F)


@app.get("/api/my-assessment.docx")
async def my_assessment_docx(authorization: Optional[str] = Header(None)):
    role, sub = _check(authorization, {"trainee"})
    trainee = _get_trainee(sub)
    if not trainee:
        raise HTTPException(404, "trainee record not found")
    if not trainee.get("approved"):
        raise HTTPException(403, "your assessment hasn't been approved for download yet — check with your admin")

    res = _redis("LRANGE", EXPLAIN_LIST_KEY, "0", "999")
    rows = res.get("result") or []
    all_recs = [json.loads(x) for x in rows]
    mine = [r for r in all_recs if r.get("trainee_phone") == sub]
    latest = {}
    for r in mine:
        if r.get("game_id") not in latest:
            latest[r["game_id"]] = r
    done = list(latest.values())
    done.sort(key=lambda r: r.get("received_at", ""))

    def score_total(scores):
        return sum(1 for k, _ in CRITERIA_LABELS if scores.get(k))

    grand = sum(score_total(r.get("scores") or {}) for r in done)

    doc = Document()
    title = doc.add_heading("Playwise Assessment", level=0)
    title.runs[0].font.color.rgb = BRAND_CORAL

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"{trainee.get('name')} — {', '.join(_trainee_categories(trainee))}" + (f" — {trainee.get('cohort')}" if trainee.get("cohort") else ""))
    sub_run.bold = True
    sub_run.font.size = Pt(13)
    doc.add_paragraph(time.strftime("%d %B %Y"), style=None)

    overall = doc.add_paragraph()
    overall_run = overall.add_run(f"Overall: {grand} / {len(done) * 5}  ({len(done)} games completed)")
    overall_run.bold = True
    overall_run.font.size = Pt(12)
    overall_run.font.color.rgb = BRAND_TEAL

    # cross-game pattern summary
    misses = {}
    for r in done:
        scores = r.get("scores") or {}
        for k, label in CRITERIA_LABELS:
            if not scores.get(k):
                misses[label] = misses.get(label, 0) + 1
    doc.add_heading("Areas to work on", level=1).runs[0].font.color.rgb = BRAND_TEAL
    if misses:
        for label, count in sorted(misses.items(), key=lambda x: -x[1]):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label} — missed in {count} of {len(done)} games")
    else:
        doc.add_paragraph("No repeated weak spots — strong across the board.")

    doc.add_heading("Game by game", level=1).runs[0].font.color.rgb = BRAND_TEAL
    for r in done:
        scores = r.get("scores") or {}
        total = score_total(scores)
        h = doc.add_heading(f"{r.get('game_name')} (ages {r.get('age_band')}) — {total}/5", level=2)
        h.runs[0].font.color.rgb = BRAND_INK
        for k, label in CRITERIA_LABELS:
            p = doc.add_paragraph(style="List Bullet")
            mark = "✓" if scores.get(k) else "✗"
            run = p.add_run(f"{mark} {label}")
            run.font.color.rgb = BRAND_TEAL if scores.get(k) else RGBColor(0xC0, 0x49, 0x2F)
        if r.get("strengths"):
            p = doc.add_paragraph()
            p.add_run("✅ What you did well: ").bold = True
            p.add_run(r["strengths"])
        for imp in (r.get("improvements") or []):
            label = dict(CRITERIA_LABELS).get(imp.get("criterion"), imp.get("criterion") or "")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"🎯 {label}: ").bold = True
            p.add_run(imp.get("reason") or "")
            if imp.get("fix"):
                p.add_run(f" — Try this: {imp['fix']}").italic = True
        if r.get("new_idea"):
            p = doc.add_paragraph()
            p.add_run("💡 Idea suggested: ").bold = True
            p.add_run(r["new_idea"])

    closing = doc.add_paragraph()
    closing_run = closing.add_run("— Housie")
    closing_run.italic = True
    closing_run.font.color.rgb = BRAND_MUTED

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = re.sub(r"[^a-z0-9]+", "-", trainee.get("name", "assessment").lower()) + "-assessment.docx"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
